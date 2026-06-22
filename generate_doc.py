import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_design_system_doc():
    doc = Document()
    
    # Page Margins & Normal Styles Setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700
    
    # XML Helpers
    def set_cell_background(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            m = OxmlElement(margin)
            m.set(qn('w:w'), str(val))
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        
    def add_heading_styled(text, level, color_rgb):
        p = doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = color_rgb
        return p

    def add_callout(text, title="IMPORTANT DESIGN RULE", color_hex="F8FAFC", border_hex="6366F1"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, color_hex)
        set_cell_margins(cell, top=120, bottom=120, left=200, right=200)
        
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>'
            f'<w:top w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run_title = p.add_run(f"[{title}] \n")
        run_title.font.bold = True
        run_title.font.size = Pt(9.5)
        run_title.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
        
        run_body = p.add_run(text)
        run_body.font.size = Pt(9)
        run_body.font.italic = True
        run_body.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    for _ in range(3): doc.add_paragraph()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t1 = p_title.add_run("🫧 WAYAX DESIGN SYSTEM")
    run_t1.font.size = Pt(28)
    run_t1.font.bold = True
    run_t1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Organised Design System Library Master Specification")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
    
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.paragraph_format.space_after = Pt(120)
    run_sep = p_sep.add_run("____________________________________________________")
    run_sep.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta1 = p_meta.add_run("SEBI COMPLIANT AI RESEARCH ADVISORY DESIGN DIRECTORY\n")
    r_meta1.font.bold = True
    r_meta1.font.size = Pt(9.5)
    r_meta1.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    r_meta2 = p_meta.add_run("Figma Community Template: 'Organising Design Systems'\n")
    r_meta2.font.size = Pt(9.5)
    r_meta2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    r_meta3 = p_meta.add_run("IN SYNC WITH ACTIVE REACT & CSS PARAMS (Version 2.1.0)")
    r_meta3.font.size = Pt(9)
    r_meta3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # FIGMA PAGE 01: WELCOME & PRINCIPLES
    # ----------------------------------------------------
    add_heading_styled("Figma Page 01: Welcome, Status & Principles", 1, RGBColor(0x0F, 0x17, 0x2A))
    
    doc.add_paragraph(
        "Welcome to the WayaX Design System single source of truth library structure. "
        "Every frosted card in our master visual vector board maps exactly to a dedicated page inside "
        "your Figma library, keeping designers, engineers, and SEBI compliance auditors completely synchronized."
    ).paragraph_format.space_after = Pt(8)
    
    add_heading_styled("A. Design Philosophy: Refractive Specularity", 2, RGBColor(0x63, 0x66, 0xF1))
    doc.add_paragraph(
        "WayaX rejects flat user interfaces. Trust is tactile. The design language acts as a physical, thick "
        "optical glass lens layered on top of numeric datasets, providing visual isolation and safety margins. "
        "Saturate and blur filters (blur-32px saturate-185%) are combined with a specular highlight edge line "
        "at the top border of every container to simulate realistic light physics."
    ).paragraph_format.space_after = Pt(10)
    
    add_heading_styled("B. Component Matrix Status Tracker", 2, RGBColor(0x63, 0x66, 0xF1))
    
    status_data = [
        ("Frosted Action Buttons", "READY", "Page 03: Component Library", ".liquid-glass-button"),
        ("Recessed Search Wells", "READY", "Page 03: Component Library", ".liquid-glass-input"),
        ("Advisory Stock Table", "READY", "Page 05: Organisms Spec", "<StockTable />"),
        ("Presets Category Row", "READY", "Page 06: Sandbox Playground", "max-h-[132px] height-locked"),
        ("FAQ Slideout Drawer", "READY", "Page 03: Component Library", "<FAQDrawer /> at top-[76px]")
    ]
    
    table_status = doc.add_table(rows=1, cols=4)
    table_status.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table_status.rows[0].cells
    hdr_titles = ["Component Name", "Figma Status", "Library Page", "React Variable / Hook"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for name, stat, page, hook in status_data:
        row_cells = table_status.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = stat
        row_cells[2].text = page
        row_cells[3].text = hook
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            if i == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
                
    doc.add_page_break()

    # ----------------------------------------------------
    # FIGMA PAGE 02: FOUNDATIONS & DESIGN TOKENS
    # ----------------------------------------------------
    add_heading_styled("Figma Page 02: Foundations & Design Tokens", 1, RGBColor(0x0F, 0x17, 0x2A))
    
    add_heading_styled("A. Sub-Pixel Color System (Camel Deprecated)", 2, RGBColor(0x63, 0x66, 0xF1))
    doc.add_paragraph(
        "Color variables are mathematically aligned. Salted Camel (#D5AD76) has been fully deprecated and removed. "
        "The primary gradient stop uses a clean 2-stop configuration running between Teal Momentum and Sky Blue Flow."
    ).paragraph_format.space_after = Pt(8)
    
    colors_data = [
        ("Teal Momentum", "#77F2E4", "Primary Brand", "Left gradient stop (0%) of primary logos"),
        ("Sky Blue Flow", "#6BB6F3", "Primary Brand", "Right gradient stop (100%) of primary logos"),
        ("Active Indigo", "#6366F1", "Theme Accent", "Approved badges, drawer focus outlines, navigation"),
        ("Success Emerald", "#10B981", "Positive Data", "BUY calls, target upside indicators, target prices"),
        ("Protective Rose", "#EF4444", "Warning Data", "SHORT calls, Stop Loss protective warnings, volatility"),
        ("Canvas Black", "#000000", "Pitch Backdrop", "Perfect dark canvas baseline"),
        ("Slate Overlay", "#0D0E12", "Frosted Surface", "Liquid glass panels fill background")
    ]
    
    table_colors = doc.add_table(rows=1, cols=4)
    table_colors.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table_colors.rows[0].cells
    hdr_titles = ["Token Name", "Hex Code", "Role Description", "Application Scope"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for name, hexcode, role, desc in colors_data:
        row_cells = table_colors.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = hexcode
        row_cells[2].text = role
        row_cells[3].text = desc
        color_rgb = hexcode.lstrip('#')
        set_cell_background(row_cells[1], color_rgb)
        
        r_c, g_c, b_c = int(color_rgb[0:2], 16), int(color_rgb[2:4], 16), int(color_rgb[4:6], 16)
        luminance = (0.299 * r_c + 0.587 * g_c + 0.114 * b_c) / 255
        text_color = RGBColor(0x00, 0x00, 0x00) if luminance > 0.5 else RGBColor(0xFF, 0xFF, 0xFF)
        
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            if idx == 1:
                run.font.bold = True
                run.font.color.rgb = text_color
                
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Typography scales
    add_heading_styled("B. Typography Hierarchy Scales", 2, RGBColor(0x63, 0x66, 0xF1))
    doc.add_paragraph(
        "We utilize Sora (geometric Display) for brand headlines, Inter for standard text copy, "
        "and JetBrains Mono for monospaced numeric data grids scannability."
    ).paragraph_format.space_after = Pt(8)
    
    typo_data = [
        ("Welcome Headings", "Sora (sans-serif)", "28px", "text-[28px] font-extrabold", "Landing dashboard header titles"),
        ("Dropdown Selectors", "Sora (sans-serif)", "20px", "text-[20px] font-extrabold", "Preset category selection titles"),
        ("Advisory Answer Body", "Inter (sans-serif)", "13px", "text-[13px] font-medium", "AI conversation response bubbles"),
        ("Grid Numeric Values", "JetBrains Mono", "12.5px", "text-[12.5px] font-bold", "Prices, target values, stop losses")
    ]
    
    table_typo = doc.add_table(rows=1, cols=5)
    table_typo.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table_typo.rows[0].cells
    hdr_titles = ["Style Name", "Font Family", "Size", "Tailwind Class", "Application Scope"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for name, font_f, size, tw_c, scope in typo_data:
        row_cells = table_typo.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = font_f
        row_cells[2].text = size
        row_cells[3].text = tw_c
        row_cells[4].text = scope
        for cell in row_cells:
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            
    doc.add_page_break()

    # ----------------------------------------------------
    # FIGMA PAGE 03: COMPONENT LIBRARY
    # ----------------------------------------------------
    add_heading_styled("Figma Page 03: Component Library (Atoms & Molecules)", 1, RGBColor(0x0F, 0x17, 0x2A))
    
    add_heading_styled("A. Button States & Interactive Matrix", 2, RGBColor(0x63, 0x66, 0xF1))
    
    btn_data = [
        ("Primary Action", "bg-indigo-600 border border-indigo-500", "bg-indigo-500 translate-y-[-1.5px]", "scale-95 bg-indigo-750", "opacity-50 cursor-not-allowed bg-slate-800"),
        ("Secondary Glass", "bg-white/[0.05] border border-white/10 backdrop-blur", "bg-white/[0.08] translate-y-[-1.5px] (sweeps .liquid-shimmer)", "scale-[0.98] bg-white/[0.12]", "opacity-40 cursor-not-allowed"),
        ("Tertiary Capsule", "bg-white/[0.02] border border-white/5", "bg-white/[0.08] text-white", "bg-white/[0.15]", "hidden")
    ]
    
    table_btn = doc.add_table(rows=1, cols=5)
    table_btn.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table_btn.rows[0].cells
    hdr_titles = ["Button Type", "Default State", "Hover State (Translate & Shimmer)", "Active State (Click)", "Disabled State"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for b_type, def_s, hov_s, act_s, dis_s in btn_data:
        row_cells = table_btn.add_row().cells
        row_cells[0].text = b_type
        row_cells[1].text = def_s
        row_cells[2].text = hov_s
        row_cells[3].text = act_s
        row_cells[4].text = dis_s
        for cell in row_cells:
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading_styled("B. Recessed Search Input Wells", 2, RGBColor(0x63, 0x66, 0xF1))
    doc.add_paragraph(
        "Search fields are designed as recessed visual wells carved into the panel surface:\n"
        "•  Default: bg-black/25, border border-white/7, inner-shadow inset 0 2px 4px rgba(0,0,0,0.35).\n"
        "•  Focus: border-white/20, bg-black/35, shadow-[inset_0_2px_8px_rgba(0,0,0,0.5),0_0_24px_rgba(255,255,255,0.06)], "
        "triggering an outline breathing pulse animation."
    ).paragraph_format.space_after = Pt(8)
    
    add_callout(
        "To prevent height jumps when dropdown questions toggle, the preset container is restricted strictly to "
        "a max-height showing exactly 3.5 elements (3 items + 30% faded scroll visual hint of the 4th item).",
        title="SCROLL PRESETS HEIGHT LAWS"
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # FIGMA PAGE 04: PATTERN GUIDELINES
    # ----------------------------------------------------
    add_heading_styled("Figma Page 04: Pattern Guidelines & Handoff", 1, RGBColor(0x0F, 0x17, 0x2A))
    
    add_heading_styled("🟢 The Emerald Green Law (Action & Upside)", 2, RGBColor(0x10, 0xB9, 0x81))
    doc.add_paragraph(
        "Green represents positive wealth-building signals and target indicators:\n"
        "•  BUY Action pills use bg-emerald-500/10 with an active pulsing status indicator dot (bg-emerald-500 animate-pulse).\n"
        "•  Target upside values are encased inside green border tags containing a TrendingUp icon.\n"
        "•  Future targets are styled in monospace bold text-emerald-500."
    ).paragraph_format.space_after = Pt(10)
    
    add_heading_styled("🔴 The Rose Red Law (Protection & Warnings)", 2, RGBColor(0xEF, 0x44, 0x44))
    doc.add_paragraph(
        "Red is reserved exclusively for risk guidance and protective stop-losses:\n"
        "•  SHORT Action pills use bg-rose-500/10 with an active pulsing status dot.\n"
        "•  Stop Loss (SL) values are styled in bold warning text-rose-400.\n"
        "•  Leverage alert warning is triggered red when Debt-to-Equity exceeds critical limits (Debt/Equity > 1.50)."
    ).paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # FIGMA PAGE 05: STOCK TABLE SPEC
    # ----------------------------------------------------
    add_heading_styled("Figma Page 05: Stock Table Organism Spec", 1, RGBColor(0x0F, 0x17, 0x2A))
    
    doc.add_paragraph(
        "To prevent layout grid shifts when rows expand, WayaX tables utilize strict pixel metric dimensions:"
    ).paragraph_format.space_after = Pt(8)
    
    grid_data = [
        ("# Index", "w-14 / Centered", "JetBrains Mono", "text-slate-500", "01, 02 index numbering padded with leading zeros"),
        ("Symbol Name", "min-w-[220px] / Left", "Sora & Mono", "text-white / bg-white/5", "Company name stacked over symbol badge in bg-white/5"),
        ("Call Action", "w-32 / Centered", "Inter Bold Caps", "Emerald / Rose", "BUY or SHORT status badges with active pulsing dot"),
        ("Trigger Zone", "Flexible / Centered", "JetBrains Mono", "text-[#6bb6f3]", "Blue monospace zone values (text-[#6bb6f3])"),
        ("Target Price", "Flexible / Centered", "JetBrains Mono", "text-emerald-500", "Success Emerald bold future targets"),
        ("Protective SL", "Flexible / Centered", "JetBrains Mono", "text-rose-400", "Protective Rose semibold risk stop guide"),
        ("Details Button", "w-32 / Centered", "Inter Bold", "Glass Backing", "Capsule button triggers secondary drawer accordion")
    ]
    
    table_grid = doc.add_table(rows=1, cols=5)
    table_grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table_grid.rows[0].cells
    hdr_titles = ["Column Name", "Grid Width / Align", "Font Family", "Color Token", "Cell Details & Context"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=150, right=150)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for name, width_align, font_family, color_tok, desc in grid_data:
        row_cells = table_grid.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = width_align
        row_cells[2].text = font_family
        row_cells[3].text = color_tok
        row_cells[4].text = desc
        for cell in row_cells:
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    doc.add_paragraph(
        "•  Row Spacing: Spaced with vertical py-4.5 (18px) and horizontal px-5 (20px).\n"
        "•  Row Hover highlights: hover:bg-white/[0.035] (expanded active background locks to bg-white/[0.015]).\n"
        "•  Accordion Drawer Detail: Click 'Analyze' opens a nested drawer spanning 9 columns. "
        "Separated into three areas: Technical Matrix, Fundamental Shield, and the italics-styled investment thesis."
    ).paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()

    # ----------------------------------------------------
    # FIGMA PAGE 06: PLAYGROUND & CHANGELOG
    # ----------------------------------------------------
    add_heading_styled("Figma Page 06: Playground Sandbox & Changelog", 1, RGBColor(0x0F, 0x17, 0x2A))
    
    add_heading_styled("A. Playground Sandbox", 2, RGBColor(0x63, 0x66, 0xF1))
    doc.add_paragraph(
        "Provides a canvas environment where designers can prototype and test components "
        "without affecting the production libraries variables structure."
    ).paragraph_format.space_after = Pt(8)
    
    add_heading_styled("B. Release Changelog", 2, RGBColor(0x63, 0x66, 0xF1))
    doc.add_paragraph(
        "•  v2.1.0: Redesigned the design system board itself to use caustics refraction circles, dotted canvas grid patterns, and specular highlights card boundaries.\n"
        "•  v2.0.4: Re-arranged the entire visual canvas into a 2x3 grid mapping exactly to the 6 standard directories of the 'Organising Design Systems' community template.\n"
        "•  v2.0.0: Deprecated Salted Camel color hex (#D5AD76) from gradient stops, updating gradientstops strictly to Teal-to-Blue configuration (#77F2E4 and #6BB6F3).\n"
        "•  v1.8.2: Integrated exact 3-layered brand logo underlay refraction blurs and golden emblem asset."
    ).paragraph_format.space_after = Pt(12)
    
    # Save the document
    output_path = "WayaX_Design_System.docx"
    doc.save(output_path)
    print(f"Design System Document successfully saved to {output_path}")

if __name__ == "__main__":
    create_design_system_doc()
